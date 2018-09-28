from docx import Document
from docx.shared import Inches

document = Document()

paragraph = document.add_paragraph("py is best language in the world");


table = document.add_table(rows=1,cols=3)

data = (
    ('用膳','100','90'),
    ('泡澡','200','180'),
    ('洗脚','300','200')
)

row = table.rows[0] #获取本张表的第一行
row.cells[0].text = '服务项目'   #本行的第一个单元格
row.cells[1].text = '费用'
row.cells[2].text = '优惠价格'

for d in data:
    new_row = table.add_row().cells  #给本张表添加一行并取得该行的所有单元格列表　　数组

    new_row[0].text = d[0]
    new_row[1].text = d[1]
    new_row[2].text = d[2]


document.add_picture('https://ss1.bdstatic.com/70cFuXSh_Q1YnxGkpoWK1HF6hhy/it/u=2639088341,2223755776&fm=27&gp=0.jpg');
document.save("test8.docx");

""""

#paragraph_before = paragraph.insert_paragraph_before("家伙");

document.add_heading("我是一个大大的标题");

document.add_heading("我又是一个大大的标题1",level=2);
document.add_heading("我又是一个大大的标题2",level=3);
document.add_heading("我又是一个大大的标题3",level=4);

document.add_page_break()
cell0 = table.cell(0,0)
cell1 = table.cell(0,1)
cell2 = table.cell(0,2)


cell0.text = '学校'
cell1.text = '校长'
cell2.text = '历史'

table.cell(1,0).text = '清华'
table.cell(1,1).text = '不晓得'
table.cell(1,2).text = '100年'

row = table.rows[2]

#document.add_paragraph("这个word文档有"+rows+"行,每行有"+cells+'个单元格')

row.cells[0].text = '北大'
row.cells[1].text = '张三'
row.cells[2].text = '200年'
data = (
    ('名字'),
    ('年龄'),
    ('技能')
)
i,j=0,0
for r in table.rows:

    for cell in r.cells:
        cell.text = data[i][j]
        j=j+１
    i=i+1

"""

